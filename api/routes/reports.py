from flask import Blueprint, request, Response, jsonify, send_file
from supabase import create_client, Client
from datetime import datetime
from dotenv import load_dotenv
import os
from models.agent_teams import create_research_team, create_writing_team, create_report_generator
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import matplotlib.pyplot as plt
import seaborn as sns
import io
import json
import logging
import uuid
from models.llm import get_llm
from langchain.schema import HumanMessage

# Load environment variables
load_dotenv()

# Initialize Blueprint
reports_bp = Blueprint('reports', __name__)
logger = logging.getLogger(__name__)

# Get Supabase credentials
SUPABASE_URL = os.getenv('SUPABASE_URL')
SUPABASE_KEY = os.getenv('SUPABASE_SERVICE_KEY')

# Validate Supabase credentials
if not SUPABASE_URL or not SUPABASE_KEY:
    raise ValueError(
        "Missing Supabase credentials. Please ensure SUPABASE_URL and SUPABASE_SERVICE_KEY "
        "are set in your environment variables."
    )

# Initialize Supabase client
try:
    supabase = create_client(SUPABASE_URL, SUPABASE_KEY)
except Exception as e:
    print(f"Error initializing Supabase client: {str(e)}")
    raise

# Create reports directory if it doesn't exist
REPORTS_DIR = 'reports'
os.makedirs(REPORTS_DIR, exist_ok=True)

# Ensure the directory has proper permissions
try:
    os.chmod(REPORTS_DIR, 0o755)  # rwxr-xr-x
except Exception as e:
    logger.warning(f"Could not set directory permissions: {str(e)}")

# Initialize the LLM
llm = get_llm()

def format_sse(data):
    """Format data as SSE."""
    return f"data: {json.dumps(data)}\n\n"

def create_chart(data, chart_type='line', title='', xlabel='', ylabel=''):
    """Create a chart using matplotlib and return it as bytes."""
    plt.figure(figsize=(10, 6))
    
    if chart_type == 'line':
        plt.plot(data['x'], data['y'])
    elif chart_type == 'bar':
        plt.bar(data['x'], data['y'])
    elif chart_type == 'pie':
        plt.pie(data['values'], labels=data['labels'])
    
    plt.title(title)
    plt.xlabel(xlabel)
    plt.ylabel(ylabel)
    
    # Save plot to bytes
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=300, bbox_inches='tight')
    plt.close()
    buf.seek(0)
    return buf

def add_chart_to_doc(doc, chart_data, chart_type='line', title='', xlabel='', ylabel=''):
    """Add a chart to the document."""
    # Create chart
    chart_buf = create_chart(chart_data, chart_type, title, xlabel, ylabel)
    
    # Add chart to document
    doc.add_picture(chart_buf, width=Inches(6))
    doc.add_paragraph()  # Add spacing

def create_docx_report(report_data, report_id):
    """Create a DOCX report with proper formatting and structure."""
    try:
        # Create a new document
        doc = Document()
        
        # Title
        title = doc.add_heading('Consulting Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        metadata = doc.add_paragraph()
        metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
        metadata.add_run(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}\n')
        metadata.add_run(f'Report ID: {report_id}')
        doc.add_paragraph()
        
        # Table of Contents
        doc.add_heading('Table of Contents', 1)
        toc = doc.add_paragraph()
        toc.add_run('1. Executive Summary\n')
        toc.add_run('2. Findings\n')
        toc.add_run('3. Analysis\n')
        toc.add_run('4. Conclusions\n')
        toc.add_run('5. Recommendations\n')
        doc.add_paragraph()
        
        # Executive Summary
        doc.add_heading('1. Executive Summary', 1)
        summary = doc.add_paragraph()
        summary.add_run(report_data['executive_summary'])
        doc.add_paragraph()
        
        # Research Findings
        doc.add_heading('2. Findings', 1)
        findings = doc.add_paragraph()
        if isinstance(report_data['research_findings'], list):
            for finding in report_data['research_findings']:
                findings.add_run(f'• {finding}\n')
        else:
            findings.add_run(report_data['research_findings'])
        doc.add_paragraph()
        
        # Analysis
        doc.add_heading('3. Analysis', 1)
        analysis = doc.add_paragraph()
        if isinstance(report_data['analysis'], dict):
            for section, content in report_data['analysis'].items():
                analysis.add_run(f'{section}:\n', bold=True)
                analysis.add_run(f'{content}\n\n')
        else:
            analysis.add_run(report_data['analysis'])
        doc.add_paragraph()
        
        # Conclusions
        doc.add_heading('4. Conclusions', 1)
        conclusions = doc.add_paragraph()
        conclusions.add_run(report_data['conclusions'])
        doc.add_paragraph()
        
        # Recommendations
        doc.add_heading('5. Recommendations', 1)
        recommendations = doc.add_paragraph()
        if isinstance(report_data['recommendations'], list):
            for rec in report_data['recommendations']:
                recommendations.add_run(f'• {rec}\n')
        else:
            recommendations.add_run(report_data['recommendations'])
        
        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run('End of Report')
        
        # Save to memory buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer
        
    except Exception as e:
        logger.error(f"Error creating DOCX report: {str(e)}")
        raise

@reports_bp.route('/reports', methods=['GET'])
async def get_reports():
    """Get all reports for the current user."""
    try:
        # Get user ID from token
        auth_header = request.headers.get('Authorization')
        if not auth_header or not auth_header.startswith('Bearer '):
            return jsonify({
                'error': 'Invalid authorization header',
                'success': False
            }), 401
            
        token = auth_header.split(' ')[1]
        user_id = await supabase.auth.get_user(token).data.id

        # Fetch reports from Supabase
        response = supabase.table('reports').select('*').eq('user_id', user_id).execute()
        
        if response.error:
            raise Exception(response.error.message)
            
        return jsonify({
            'success': True,
            'reports': response.data
        })
    except Exception as e:
        return jsonify({
            'error': str(e),
            'success': False
        }), 500

@reports_bp.route('/reports/<report_id>', methods=['GET'])
async def get_report(report_id):
    """Get a specific report by ID."""
    try:
        # Get user ID from token
        auth_header = request.headers.get('Authorization')
        if not auth_header or not auth_header.startswith('Bearer '):
            return jsonify({
                'error': 'Invalid authorization header',
                'success': False
            }), 401
            
        token = auth_header.split(' ')[1]
        user_id = await supabase.auth.get_user(token).data.id

        # Fetch report from Supabase
        response = supabase.table('reports').select('*').eq('id', report_id).eq('user_id', user_id).single().execute()
        
        if response.error:
            raise Exception(response.error.message)
            
        return jsonify({
            'success': True,
            'report': response.data
        })
    except Exception as e:
        return jsonify({
            'error': str(e),
            'success': False
        }), 500

@reports_bp.route('/reports/<report_id>', methods=['DELETE'])
async def delete_report(report_id):
    """Delete a specific report."""
    try:
        # Get user ID from token
        auth_header = request.headers.get('Authorization')
        if not auth_header or not auth_header.startswith('Bearer '):
            return jsonify({
                'error': 'Invalid authorization header',
                'success': False
            }), 401
            
        token = auth_header.split(' ')[1]
        user_id = await supabase.auth.get_user(token).data.id

        # Delete report from Supabase
        response = supabase.table('reports').delete().eq('id', report_id).eq('user_id', user_id).execute()
        
        if response.error:
            raise Exception(response.error.message)
            
        return jsonify({
            'success': True,
            'message': 'Report deleted successfully'
        })
    except Exception as e:
        return jsonify({
            'error': str(e),
            'success': False
        }), 500

@reports_bp.route('/generate', methods=['POST'])
def generate_report():
    """Handle report generation requests."""
    try:
        data = request.get_json()
        query = data.get('query')
        
        if not query:
            return jsonify({'error': 'No query provided'}), 400
        
        report_id = str(uuid.uuid4())
        
        def generate():
            # Initialize report generator
            report_generator = create_report_generator(llm)
            
            # Initialize state with the query
            initial_state = {
                "messages": [
                    HumanMessage(content=query, name="user")
                ],
                "status": "started",
                "phase": "research"
            }
            
            # Research phase
            yield format_sse({
                'type': 'status',
                'agent': 'researcher',
                'content': 'Gathering information...'
            })
            yield format_sse({
                'type': 'progress',
                'progress': 20
            })
            
            # Invoke report generator
            results = report_generator.invoke(initial_state)
            
            # Extract report content from results
            report_content = results['messages'][-1].content
            
            # Stream the content to frontend
            yield format_sse({
                'type': 'report',
                'reportId': report_id,
                'content': report_content
            })
            
            yield format_sse({
                'type': 'progress',
                'progress': 100
            })
        
        return Response(generate(), mimetype='text/event-stream')
    
    except Exception as e:
        logger.error(f"Error in generate_report: {str(e)}")
        return jsonify({'error': str(e)}), 500

@reports_bp.route('/<report_id>/download', methods=['GET'])
def download_report(report_id):
    """Handle report download requests."""
    try:
        # Get report content from the request
        report_content = request.args.get('content')
        if not report_content:
            return jsonify({'error': 'No content provided'}), 400

        # Create report data structure
        report_data = {
            'executive_summary': report_content,
            'research_findings': report_content,
            'analysis': report_content,
            'conclusions': report_content,
            'recommendations': report_content
        }

        # Create DOCX document
        doc_buffer = create_docx_report(report_data, report_id)
        
        # Return the document for download
        return Response(
            doc_buffer.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={
                'Content-Disposition': f'attachment; filename=report_{report_id}.docx'
            }
        )
        
    except Exception as e:
        logger.error(f"Error downloading report: {str(e)}")
        return jsonify({'error': str(e)}), 500

@reports_bp.route('/generate', methods=['OPTIONS'])
def generate_report_options():
    """Handle CORS preflight requests."""
    response = Response()
    response.headers.add('Access-Control-Allow-Origin', '*')
    response.headers.add('Access-Control-Allow-Headers', 'Content-Type,Authorization')
    response.headers.add('Access-Control-Allow-Methods', 'POST,OPTIONS')
    return response 